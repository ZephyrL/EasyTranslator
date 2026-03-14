import docx

from src.models.document import Document, Paragraph


def parse_docx(file_path: str) -> Document:
    """Parse a .docx file into paragraphs using python-docx."""
    doc = docx.Document(file_path)
    paragraphs = []
    idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(Paragraph(index=idx, text=text))
            idx += 1

    return Document(file_path=file_path, paragraphs=paragraphs)
